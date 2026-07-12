#!/usr/bin/env python3
"""更新专利交底书 — 加入仿真联调内容"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = os.path.expanduser('~/lerobot-smolvla-lew/docs/patents/Z-MAX-专利交底书-实用新型.docx')
doc = Document(doc_path)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

# 在摘要前插入仿真联调章节
# 找到"七、摘要"并在此之前插入
for i, p in enumerate(doc.paragraphs):
    if '七、摘要' in p.text:
        # 在摘要前插入新章节
        insert_idx = i
        
        # 找到插入点的元素(需要操作XML)
        break

# 直接在末尾添加(附加到摘要后面)
add_heading_styled('八、仿真联调系统', level=1)
doc.add_paragraph(
    '本实用新型还包含一套完整的仿真联调系统，用于在离线或开发环境中'
    '验证视觉语言动作模型的推理效果，以及与机器人控制系统的集成。'
)

add_heading_styled('8.1 仿真架构', level=2)
doc.add_paragraph(
    '仿真联调系统采用Client-Server架构，基于WebSocket JSON协议通信：\n'
    '（a）Client端（运行于Mac M1端侧）：模拟五种传感器——6轴关节传感器、六维力/扭矩传感器'
    '（1kHz仿真采样）、RealSense D435相机（640×480 RGB-D）、电动夹爪、4×4触觉阵列。'
    'Client以30Hz频率将传感器数据打包为JSON消息发送至Server。\n'
    '（b）Server端（运行于GPU服务器/WSL2）：加载SmolVLA或ACT模型，'
    '接收传感器数据后执行推理，将预测的6轴关节位置和夹爪开度作为动作指令返回Client。\n'
    '（c）通信协议：11个/sim/命名空间的话题，传感器包1061B，动作包287B，30Hz仅需0.25Mbps带宽。'
)

add_heading_styled('8.2 双模型分流策略', level=2)
doc.add_paragraph(
    '仿真系统支持双模型分级部署：\n'
    '（1）ACT 51.6M模型：用于高速重复操作仿真，推理延迟<1ms（Mac MPS），帧率>2000FPS，'
    '适合基础动作验证和高速产线模拟。\n'
    '（2）SmolVLA 450M模型：用于端到端感知-决策-执行仿真，推理延迟约215ms（RTX4060），'
    '具备视觉场景理解、自然语言指令解析和精细动作生成能力。\n'
    '（3）SmolVLA-LEW Mini 263K模型：自训练轻量版本，推理延迟<1ms，适合嵌入式端侧部署验证。'
)

add_heading_styled('8.3 集成测试验证', level=2)
doc.add_paragraph(
    '仿真系统通过了5项集成测试：协议编解码验证（action 289B/sensor 1061B）、'
    '5类传感器模拟器输出验证、Client独立模式100包连续发布（16,033Hz纯本地速率）、'
    '30Hz持续通信带宽基准（0.25Mbps）和11话题命名兼容性验证。'
    'Client-Server往返延迟在局域网环境下<5ms。'
)

# 保存
doc.save(doc_path)
print(f"✅ 专利交底书已更新: {doc_path}")
print(f"   新增: 第八章 仿真联调系统 (8.1架构 + 8.2双模型 + 8.3集成测试)")
