#!/usr/bin/env python3
"""生成 Z-MAX 专利交底书（实用新型）"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# === 页面设置 ===
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

# === 封面 ===
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('专 利 交 底 书')
run.bold = True
run.font.size = Pt(22)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('（实用新型）')
run.font.size = Pt(14)

doc.add_paragraph()

info = [
    ('发明名称', 'Z-MAX 多模态视觉语言动作模型驱动的\n光模块自主插拔机器人系统'),
    ('申请人', '智蜂创元（ZFCY）'),
    ('技术领域', '具身智能机器人 · 光模块自动化测试与装配'),
    ('发明人', '待填写'),
    ('申请日期', '2026年07月'),
]
for label, value in info:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}：')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(value)

doc.add_page_break()

# === 一、技术领域 ===
add_heading_styled('一、技术领域')
doc.add_paragraph(
    '本实用新型涉及具身智能机器人技术领域，具体涉及一种基于多模态视觉语言动作模型（VLA）'
    '和类脑三层解耦架构的光模块自主插拔机器人系统，适用于高速光模块（400G/800G/1.6T）'
    '在测试与装配环节的全自主精细操作。'
)

# === 二、背景技术 ===
add_heading_styled('二、背景技术')
doc.add_paragraph(
    '高速光模块是AI算力基础设施的核心硬件载体。随着800G/1.6T光模块需求爆发，'
    '光模块测试与装配环节面临严峻挑战：'
)
problems = [
    '人工插拔力度不可控，导致光模块接口损伤，良率波动大（人工插拔良率通常低于98%）；',
    '光模块型号多样（400G/800G/1.6T等多种封装），换型需要人工重新调试，无法柔性生产；',
    '密集测试仓空间狭小，人工操作困难，且存在静电损伤风险；',
    '现有自动化专机只能适应单一型号，换型成本高、周期长，无法满足多品种小批量需求；',
    '传统力控方案响应速度不足（<100Hz），无法实现微牛级精细力控保护。',
]
for prob in problems:
    doc.add_paragraph(prob, style='List Bullet')

doc.add_paragraph(
    '因此，亟需一种具备多模态感知、自主决策、精细力控能力的具身智能机器人系统，'
    '实现光模块插拔工序的全自主、高精度、柔性化操作。'
)

# === 三、发明内容 ===
add_heading_styled('三、发明内容')

add_heading_styled('3.1 要解决的技术问题', level=2)
doc.add_paragraph(
    '本实用新型要解决的技术问题是：克服现有光模块插拔自动化方案中型号适应性差、'
    '力控精度不足、换型成本高的缺陷，提供一种基于VLA模型和类脑三层架构的自主插拔系统。'
)

add_heading_styled('3.2 技术方案', level=2)
doc.add_paragraph(
    '本实用新型提供一种多模态视觉语言动作模型驱动的光模块自主插拔机器人系统，包括以下四个核心子系统：'
)

solutions = [
    ('（一）感知层（Sys-2 数据系统）',
     '包括3D深度相机（Gemini 335L）、腕部视觉相机、六维力/扭矩传感器（采样频率>10kHz）、'
     '360°鱼眼相机阵列和对角双激光雷达。感知层负责采集光模块的视觉特征（型号、位姿、插槽状态）、'
     '力学特征（接触力、插入阻力）和环境特征（障碍物、工位状态），形成多模态感知数据流。'),
    ('（二）认知层（Sys-11 训练系统）',
     '包括基于SmolVLA架构的视觉语言动作模型。该模型采用类脑双通路设计：'
     '（a）腹侧通路——冻结的SmolVLM2视觉语言模型（约350M参数），负责场景理解、光模块型号识别和自然语言指令解析；'
     '（b）背侧通路——可训练的Expert动作决策网络（约98M参数），通过交叉注意力机制读取VLM特征，'
     '生成包含6自由度位姿（x, y, z, roll, pitch, yaw）和夹爪开度的动作指令。'
     'Expert仅占全量参数的22%，支持冻结VLM仅微调Expert的快速领域迁移训练模式。'),
    ('（三）执行层（Sys-12 推理系统）',
     '包括>10kHz力控闭环执行模块，通过gRPC协议与认知层通信。'
     '执行层在光模块插入过程中实时监测六维力数据，当检测到异常阻力时自动触发保护策略：'
     '回退至安全位姿→微调末端姿态（±2°）→降低插入速度（50%）→重新尝试，最多重试3次。'
     '执行层支持旁路验证模式——推理在独立gRPC通道运行，不干扰主控PLC的实时控制回路。'),
    ('（四）OTA升级子系统',
     '支持从L2分段式自动化到L3条件自动化再到L4高度自动化的软件在线升级路径。'
     '硬件平台（Z700轮式双臂机器人，搭载NVIDIA AGX Orin边缘计算平台）一步到位，'
     '仅通过软件升级即可实现：L2（Sys-0人工编排引擎）→L3（Sys-1 SmolVLA端到端引擎）'
     '→L4（Sys-11潜空间压缩+精细感知引擎）的能力跃迁，无需更换任何硬件。'),
]
for title, desc in solutions:
    p = doc.add_paragraph()
    run = p.add_run(title + '：')
    run.bold = True
    p.add_run(desc)

add_heading_styled('3.3 有益效果', level=2)
doc.add_paragraph('与现有技术相比，本实用新型具有以下有益效果：')
effects = [
    '插拔精度达到±0.02mm，关键工序良率≥99.2%（L2基线版），最高可达≥99.9%（L4旗舰版），比人工插拔提升2个百分点以上；',
    '力控响应频率>10kHz，比传统PLC力控方案（<100Hz）提升1000倍，实现微牛级精细力保护，有效避免光模块接口损伤；',
    'Expert动作决策网络仅约100M参数，新光模块型号适配仅需微调Expert（训练速度提升4倍，GPU显存需求仅~2GB），实现小时级快速换型；',
    '硬件一次投入，通过OTA软件升级实现L2→L3→L4三级能力跃迁，无需更换设备，保护用户投资；',
    '单机可替代3名操作工，ROI投资回收周期14~22个月，具备明确的经济效益。',
]
for i, eff in enumerate(effects, 1):
    doc.add_paragraph(f'（{i}）{eff}')

doc.add_page_break()

# === 四、附图说明 ===
add_heading_styled('四、附图说明')
figures = [
    '图1：Z-MAX系统三层解耦架构示意图（感知层→认知层→执行层）；',
    '图2：SmolVLA类脑双通路模型架构图（腹侧VLM通路+背侧Expert通路+交叉注意力）；',
    '图3：力控自适应插拔闭环控制流程图（监测→判断→回退→微调→重试）；',
    '图4：OTA软件升级路径图（L2 Sys-0→L3 Sys-1→L4 Sys-11）；',
    '图5：Z700轮式双臂机器人硬件结构图（底盘+双臂+传感器+算力平台）。',
]
for f in figures:
    doc.add_paragraph(f)

# === 五、具体实施方式 ===
add_heading_styled('五、具体实施方式')

add_heading_styled('5.1 实施例1：Z700 F基线版（L2分段式自动化）', level=2)
doc.add_paragraph(
    '本实施例中，硬件平台采用Z700轮式双臂机器人，搭载NVIDIA AGX Orin边缘计算平台（2048 CUDA核心，64 Tensor核心）。'
    '机械臂为XMS5-R800六轴力控机械臂（工作半径800mm，重复定位精度±0.02mm），'
    '末端集成电动夹爪（DH-Robotics）和六维力/扭矩传感器（TS-T-15，采样率>10kHz）。'
    '视觉系统包括头部3D深度相机（Gemini 335L，分辨率1280×800）和腕部RGB相机。'
)
doc.add_paragraph('操作流程如下：')
l2_steps = [
    '人工编排：操作员通过Z-MAX Studio图形化界面，以拖拽方式定义插拔工序流程，包括：取料（左臂从料盘吸取光模块）→扫码（腕部相机识别模块二维码）→对准（移动到插槽上方预设位姿）→插入（以预设速度和力阈值向下插入）→AOI验证（腕部相机拍摄插后状态）→下料（右臂将完成模块放入成品盒）。',
    '原子功能执行：系统调用预设的标准原子功能库（基于ROS2 Service接口），按照人工编排的顺序依次执行各原子功能。每个原子功能为不可分割的最小操作单元。',
    '力控保护：六维力传感器以>10kHz频率实时监测插入过程中的Z向力。当Z向力超过预设阈值（通常设为5N）时，系统立即触发急停，机械臂停止运动并保持当前位置。',
    'AOI验证：腕部相机拍摄插后状态图像（分辨率1920×1080），通过预设的模板匹配算法（归一化相关系数>0.95判定合格），判断光模块是否完全插入到位。',
    '异常处理：如AOI判断不合格，系统通过三色塔灯（红-黄-绿）和声光报警器通知操作员，等待人工介入处理。',
]
for s in l2_steps:
    doc.add_paragraph(s, style='List Number')

add_heading_styled('5.2 实施例2：L3增强版（条件自动化）', level=2)
doc.add_paragraph(
    '本实施例在L2硬件基础上，通过OTA升级加载SmolVLA端到端模型（Sys-1引擎），无需更换任何硬件。'
    'SmolVLA模型采用SmolVLM2-500M作为视觉语言编码器（冻结），DiT-B（262M参数）作为动作头（可训练）。'
    '系统新增以下能力：'
)
l3_steps = [
    '自动视觉识别：系统启动后，头部3D相机自动扫描工位区域，VLM实时识别光模块型号（支持400G QSFP-DD、800G OSFP、1.6T等多种封装），识别准确率>99%，自动加载对应的操作配方参数（插入深度、力阈值、速度曲线）。',
    '视觉引导对准：SmolVLA模型以RGB图像和自然语言指令（如"将400G光模块插入左侧第3个插槽"）为输入，输出6自由度相对位姿（Δx, Δy, Δz, Δroll, Δpitch, Δyaw），引导夹爪精确对准插槽中心，对准误差<0.1mm。',
    '力控自适应：插入过程中，执行层以1kHz频率读取六维力数据并反馈给认知层。认知层根据力反馈动态调整动作指令——当检测到侧向力增大时自动微调末端姿态（±2°范围），当Z向力异常增大时降低插入速度。',
    '异常自恢复：当插入失败（AOI判定不合格或力超限）时，系统自动执行恢复流程：回退至安全高度（Z轴+20mm）→微调X/Y位置（±0.5mm随机偏移）→降低插入速度至50%→重新尝试，最多重试3次。3次均失败则报警并标记该工位。',
    '换线自动切换：当产线切换光模块型号时，操作员仅需将新型号样品置于视觉识别区，系统自动完成型号录入、配方生成和模型适配（Expert微调约15分钟），无需人工编程。',
]
for s in l3_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# === 六、权利要求书 ===
add_heading_styled('六、权利要求书')

add_heading_styled('6.1 独立权利要求', level=2)

p = doc.add_paragraph()
run = p.add_run('权利要求1：')
run.bold = True
p.add_run(
    '一种多模态视觉语言动作模型驱动的光模块自主插拔机器人系统，其特征在于，包括：\n'
    '（a）感知层，包括3D深度相机、腕部视觉相机、六维力/扭矩传感器，用于采集光模块的多模态感知数据；\n'
    '（b）认知层，包括基于SmolVLA架构的视觉语言动作模型，该模型采用冻结的视觉语言模型VLM编码器作为腹侧通路，'
    '与可训练的Expert动作决策网络作为背侧通路，两者通过交叉注意力机制连接，'
    'VLM负责场景理解和型号识别，Expert负责生成包含位姿和夹爪状态的动作指令；\n'
    '（c）执行层，包括采样频率不低于10kHz的关节力控闭环模块，通过gRPC协议与认知层通信，'
    '实现力控自适应插拔和异常自恢复；\n'
    '（d）OTA升级子系统，支持在不更换硬件的情况下，通过软件升级实现从分段式自动化到端到端自主操作的至少三个等级的迭代。'
)

add_heading_styled('6.2 从属权利要求', level=2)

dep_claims = [
    ('权利要求2', '根据权利要求1所述的系统，其特征在于，所述Expert动作决策网络参数量不超过全量模型参数的25%，支持冻结VLM仅微调Expert的快速领域迁移训练模式，训练速度相比全量微调提升至少4倍。'),
    ('权利要求3', '根据权利要求1所述的系统，其特征在于，所述感知层进一步包括360°鱼眼相机阵列和对角双激光雷达，用于实现全方位的环境感知和障碍物检测。'),
    ('权利要求4', '根据权利要求1所述的系统，其特征在于，所述执行层支持旁路验证模式，推理在独立gRPC通道运行，发送端和接收端与主控PLC控制回路物理隔离，不干扰实时控制。'),
    ('权利要求5', '根据权利要求1所述的系统，其特征在于，所述OTA升级子系统支持的至少三个自动化等级为：L2分段式自动化（人工编排+原子功能库）、L3条件自动化（端到端VLA模型+视觉引导）和L4高度自动化（潜空间压缩+精细感知+AI主动安全）。'),
    ('权利要求6', '根据权利要求1所述的系统，其特征在于，所述力控闭环模块在检测到插入阻力异常时，自动执行预设的多级保护策略，包括：第一级——回退至安全位姿；第二级——微调末端姿态；第三级——降低插入速度后重新尝试；最多执行3次自动重试。'),
    ('权利要求7', '根据权利要求1所述的系统，其特征在于，所述六维力/扭矩传感器的采样频率不低于10kHz，力控响应延迟不超过100微秒。'),
    ('权利要求8', '根据权利要求1所述的系统，其特征在于，还包括AOI自动光学检测模块，通过腕部相机拍摄插后状态图像，利用视觉模板匹配算法自动判定插拔是否合格，匹配阈值为归一化相关系数≥0.95。'),
]
for num, text in dep_claims:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}：')
    run.bold = True
    p.add_run(text)

doc.add_page_break()

# === 七、摘要 ===
add_heading_styled('七、摘要')
doc.add_paragraph(
    '本实用新型公开了一种多模态视觉语言动作模型驱动的光模块自主插拔机器人系统，'
    '包括感知层、认知层、执行层和OTA升级子系统。感知层通过3D相机、腕部相机和>10kHz力传感器'
    '采集视觉和力学多模态数据；认知层采用SmolVLA类脑双通路架构——冻结的VLM负责场景理解和型号识别，'
    '可训练的Expert（仅100M参数）负责动作决策，两者通过交叉注意力协同；'
    '执行层通过>10kHz力控闭环实现自适应插拔和三级异常自恢复；'
    'OTA子系统支持L2→L3→L4三级软件升级，硬件一步到位。'
    '本实用新型解决了光模块插拔工序中型号适应性差、力控精度不足、换型成本高的行业痛点，'
    '关键工序良率可达99.9%，单机替代3名操作工，填补了光模块柔性自主插拔领域的技术空白。'
)

# === 保存 ===
output_dir = os.path.expanduser('~/lerobot-smolvla-lew/docs/patents')
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'Z-MAX-专利交底书-实用新型.docx')
doc.save(output_path)
print(f'✅ 专利交底书已保存: {output_path}')
print(f'   文件大小: {os.path.getsize(output_path):,} bytes')
