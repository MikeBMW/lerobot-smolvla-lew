#!/usr/bin/env python3
"""生成Z-MAX专利交底书(实用新型) - docx格式"""

import os, sys
sys.path.insert(0, '/home/xspace/miniconda3/envs/lerobot/lib/python3.12/site-packages')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.makedirs("/home/xspace/lerobot-smolvla-lew/docs/patents", exist_ok=True)

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(doc, text, bold=False, size=12, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.5
    return p

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph()

add_para(doc, "专 利 交 底 书", bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

table = doc.add_table(rows=7, cols=2, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
info = [
    ("发明名称", "一种基于多模态视觉-语言-动作模型的具身机器人精细操作控制系统"),
    ("申请人", "智蜂创元（ZFCY）"),
    ("发明人", "（待填写）"),
    ("技术领域", "机器人控制、人工智能、多模态模型"),
    ("专利类型", "实用新型"),
    ("交底日期", "2026年7月"),
    ("机密等级", "内部机密 · 专利申请前禁止公开"),
]
for i, (k, v) in enumerate(info):
    cell_k = table.cell(i, 0)
    cell_v = table.cell(i, 1)
    cell_k.text = k
    cell_v.text = v
    set_cell_shading(cell_k, 'E8E8E8')
    for p in cell_k.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)

doc.add_page_break()

# ===== 一、技术领域 =====
add_heading_styled(doc, "一、技术领域", level=1)
add_para(doc, "本实用新型涉及机器人智能控制技术领域，具体涉及一种基于多模态视觉-语言-动作（VLA）模型的具身机器人精细操作控制系统，尤其适用于光模块、半导体等高精度电子元器件的自动化插拔、装配与测试场景。")

# ===== 二、背景技术 =====
add_heading_styled(doc, "二、背景技术", level=1)
add_heading_styled(doc, "2.1 现有技术问题", level=2)
add_para(doc, "在光模块制造领域，高速光模块（800G/1.6T）的测试与装配环节长期依赖人工操作，存在以下痛点：")
problems = [
    "（1）人工插拔力度不可控：人手操作力度不一致，导致光模块接口磨损、良率受损，关键工序良率通常仅90%~95%；",
    "（2）换型耗时无法柔性生产：不同型号（100G/400G/800G）需人工更换工装夹具，单次换型耗时30分钟以上；",
    "（3）密集测试仓人工作业困难：测试仓空间狭小、电磁环境复杂，人工作业效率低且存在安全隐患；",
    "（4）传统自动化专机柔性差：固定式自动化设备仅适配单一型号，无法应对多品种小批量的生产需求。",
]
for prob in problems:
    add_para(doc, prob)

add_heading_styled(doc, "2.2 现有技术方案不足", level=2)
add_para(doc, "目前工业机器人控制主要采用以下方案：")
add_para(doc, "（1）示教编程方案：需人工逐点示教，编程复杂、缺乏自适应能力，无法应对来料位姿随机变化；")
add_para(doc, "（2）2D视觉引导方案：传统2D视觉仅能处理平面定位，无法感知深度和三维姿态，插拔对孔精度不足；")
add_para(doc, "（3）力控独立方案：力传感器与视觉分离控制，缺乏多模态融合能力，异常工况下缺乏智能决策支持；")
add_para(doc, "（4）单臂机器人方案：单臂需串行完成取料和插拔，节拍时间长（>60秒/个），无法匹配高速产线需求。")

# ===== 三、发明内容 =====
add_heading_styled(doc, "三、发明内容", level=1)

add_heading_styled(doc, "3.1 要解决的技术问题", level=2)
add_para(doc, "本实用新型旨在提供一种高精度、强自适应、多模态融合的具身机器人精细操作控制系统，解决以下技术问题：")
add_para(doc, "（1）如何在来料位姿随机变化的情况下，实现±0.02mm的插拔对准精度；")
add_para(doc, "（2）如何实现视觉、力觉、触觉多模态信息的实时融合，支撑智能决策；")
add_para(doc, "（3）如何在不更换硬件的前提下，通过软件升级实现L2→L3→L4的自动化等级跃迁；")
add_para(doc, "（4）如何通过双臂协同将单次节拍从>60秒压缩至<25秒。")
add_para(doc, "（5）如何实现本地轻量模型与云端大模型的统一调度，兼顾实时响应与智能决策。")

add_heading_styled(doc, "3.2 技术方案", level=2)

add_heading_styled(doc, "3.2.1 系统总体架构", level=3)
add_para(doc, "本实用新型提供一种基于多模态VLA模型的具身机器人精细操作控制系统，包括：")
components = [
    "（a）感知层：包括3D深度相机、腕部视觉模块、六维力/力矩传感器、触觉传感器阵列；",
    "（b）模型层：包括多模态视觉-语言-动作模型（SmolVLA-LEW），该模型由视觉编码器、语言编码器、动作解码器和世界模型组成；",
    "（c）控制层：包括双臂协同控制器、力控闭环控制器（带宽>1kHz）、实时安全监控器；",
    "（d）执行层：包括轮式移动底盘、双臂力控机械臂、电动夹爪/吸盘末端执行器。",
]
for comp in components:
    add_para(doc, comp)

add_heading_styled(doc, "3.2.2 多模态VLA模型结构", level=3)
add_para(doc, "所述多模态VLA模型具有以下创新特征：")
add_para(doc, "（1）视觉-语言-动作端到端架构：视觉编码器（SmolVLM2-500M）提取场景特征，语言编码器理解任务指令，动作解码器（DiT-B扩散Transformer）生成连续动作序列；")
add_para(doc, "（2）世界模型（LeWorldModel）辅助：内嵌可学习的世界模型，预测动作执行后的环境状态变化，用于在线规划和安全校验；")
add_para(doc, "（3）类脑四阶段迭代：模拟人脑认知过程，分感知→理解→决策→执行四阶段闭环，每阶段输出作为下阶段输入，确保决策可解释性和可追溯性。")

add_heading_styled(doc, "3.2.3 双臂协同控制方法", level=3)
add_para(doc, "所述双臂协同控制器采用以下方法：")
add_para(doc, "（1）左臂负责取料工序：通过视觉识别无序来料→吸取/抓取→调整姿态→扫码识别→放置中转区；")
add_para(doc, "（2）右臂负责插拔工序：从中转区取料→视觉引导对准测试座→力控柔顺插入→等待测试→拔出→AOI检测→分类入盒；")
add_para(doc, "（3）双臂异步并行：左臂取料与右臂插拔并行执行，通过中转区解耦，实现单次节拍<25秒。")

add_heading_styled(doc, "3.2.4 力控闭环系统", level=3)
add_para(doc, "所述力控闭环控制器具有以下特征：")
add_para(doc, "（1）关节力控闭环频率>1kHz，实现毫秒级力反馈响应；")
add_para(doc, "（2）多模态力安全监测：同时采集关节力矩、末端六维力、指尖触觉信号，三路冗余校验；")
add_para(doc, "（3）自适应力控策略：根据插入阶段的阻力变化自动调节插入力度，保护光模块金手指不受损伤。")

add_heading_styled(doc, "3.2.5 软件定义自动化等级", level=3)
add_para(doc, "其特征在于，所述系统通过纯软件升级即可实现自动化等级跃迁，无需更换硬件：")
add_para(doc, "（1）L2基线版（Sys-0）：人工流程编排+标准原子功能库+分段验证，关键工序良率≥99.2%；")
add_para(doc, "（2）L3增强版（Sys-1）：多模态端到端模型+自主识别+闭环工作+异常自恢复，全工序良率≥99.5%；")
add_para(doc, "（3）L4旗舰版（Sys-11）：VLA模型+精细感知+场景引导+AI主动安全，全工序良率≥99.9%，7×24无人值守。")

# ===== 四、附图说明 =====
add_heading_styled(doc, "四、附图说明", level=1)
add_para(doc, "（附图待补充，建议包含以下图示：）")
add_para(doc, "图1：系统总体架构框图（感知层→模型层→控制层→执行层）")
add_para(doc, "图2：多模态VLA模型结构示意图（视觉编码器+语言编码器+动作解码器+世界模型）")
add_para(doc, "图3：双臂协同工作流程图（左臂取料→中转区→右臂插拔→AOI检测→分类）")
add_para(doc, "图4：力控闭环系统框图（关节力矩→末端力→触觉→三路冗余校验）")
add_para(doc, "图5：类脑四阶段迭代流程图（感知→理解→决策→执行→反馈）")
add_para(doc, "图6：软件升级路径示意图（L2→OTA→L3→OTA→L4）")

# ===== 五、具体实施方式 =====
add_heading_styled(doc, "五、具体实施方式", level=1)
add_heading_styled(doc, "5.1 实施例1：光模块EVB测试插拔", level=2)
add_para(doc, "以800G OSFP光模块的EVB测试插拔为例，系统工作流程如下：")
steps = [
    "步骤1（感知）：3D深度相机扫描来料托盘，腕部视觉模块精确定位光模块三维位姿（x, y, z, roll, pitch, yaw）；",
    "步骤2（理解）：VLA模型的语言编码器理解任务指令\"将800G光模块插入测试座1号位\"，视觉编码器提取场景特征，生成任务嵌入向量；",
    "步骤3（决策）：动作解码器基于任务嵌入生成双臂协同动作序列，世界模型预测动作效果并校验安全性；",
    "步骤4（执行）：左臂以吸盘吸取光模块，调整姿态后扫码确认型号，放置中转区；右臂从中转区夹取，力控闭环引导插入测试座，插入力<2N；",
    "步骤5（验证）：测试完成后拔出，腕部视觉进行AOI检测金手指状态，分类放入合格/不合格料盒。",
]
for step in steps:
    add_para(doc, step)

add_heading_styled(doc, "5.2 实施例2：产线换型自适应", level=2)
add_para(doc, "当产线从800G型号切换到400G型号时：")
add_para(doc, "（1）操作员扫码输入新型号信息，系统自动加载对应配方参数；")
add_para(doc, "（2）VLA模型利用语言编码器理解新型号规格（封装尺寸、引脚间距等），无需重新训练；")
add_para(doc, "（3）末端执行器自动调节夹爪开口宽度，适配不同封装尺寸；")
add_para(doc, "（4）换型总时间<60秒，相比传统专机换型（30分钟+）效率提升30倍以上。")

# ===== 六、技术效果 =====
add_heading_styled(doc, "六、技术效果", level=1)
add_para(doc, "本实用新型相比现有技术具有以下有益效果：")
effects = [
    ("1. 超高精度", "多模态VLA模型+力控闭环实现±0.02mm插拔精度，关键工序良率99%+，远超人工操作水平。"),
    ("2. 强自适应能力", "VLA模型可零样本理解新型号规格（通过语言编码器），换型无需重新编程，支持多品种小批量柔性生产。"),
    ("3. 软件定义升级", "同一硬件平台通过OTA软件升级实现L2→L3→L4自动化等级跃迁，保护客户硬件投资。"),
    ("4. 双臂高效协同", "双臂异步并行实现单次节拍<25秒，满足光模块产线节拍需求。"),
    ("5. 多模态安全冗余", "关节力矩+末端力+触觉三路力信号冗余校验，>1kHz闭环带宽确保人机协作安全。"),
    ("6. 类脑可解释决策", "四阶段迭代架构确保每步决策可追溯、可审计，满足工业场景可靠性要求。"),
]
for title, desc in effects:
    p = doc.add_paragraph()
    run_t = p.add_run(title + "：")
    run_t.bold = True
    run_t.font.size = Pt(12)
    run_d = p.add_run(desc)
    run_d.font.size = Pt(12)

# ===== 六（续）、硬件三档方案 =====
add_heading_styled(doc, "六（续）、硬件三档方案与仿真联调", level=1)

add_heading_styled(doc, "1. L2/L3/L4 三档算力平台", level=3)
add_para(doc, "所述系统支持三级硬件平台，通过更换AI域控模块实现算力升级：")

# 硬件表格
t_hw = doc.add_table(rows=4, cols=5, style='Table Grid')
for i, hdr in enumerate(["等级", "AI域控", "算力", "内存", "适用场景"]):
    c = t_hw.cell(0, i); c.text = hdr
    set_cell_shading(c, '00d4aa')
    for pp in c.paragraphs:
        for r in pp.runs: r.bold = True; r.font.color.rgb = RGBColor(0x06,0x08,0x0d); r.font.size = Pt(9)
hw_data = [
    ["L2 基线版", "Jetson Orin Nano 8GB", "40 TOPS INT8", "8GB LPDDR5", "单工位重复插拔"],
    ["L3 增强版", "Jetson Orin AGX 32GB", "275 TOPS INT8", "32GB LPDDR5", "多工位+SmolVLA端到端"],
    ["L4 旗舰版", "Orin AGX+RK3588", "275+6 TOPS", "32GB+8GB", "7×24全自主·多模型并行"],
]
for ri, row in enumerate(hw_data):
    for ci, val in enumerate(row):
        c = t_hw.cell(ri+1, ci); c.text = val
        if ri % 2 == 0: set_cell_shading(c, '0d1117')
        else: set_cell_shading(c, '080c14')
        for pp in c.paragraphs:
            for r in pp.runs: r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xc8,0xd1,0xd9)
doc.add_paragraph()

add_heading_styled(doc, "2. Client/Server 仿真联调架构", level=3)
add_para(doc, "所述系统采用Client/Server分离架构进行仿真联调：")
add_para(doc, "（a）Client端（ROS2节点）：部署于Mac/Linux主机，负责摄像头图像采集、力/触觉传感器信号模拟、接收并执行动作指令；")
add_para(doc, "（b）Server端（gRPC推理服务）：部署于WSL/GPU服务器，运行SmolVLA/ACT/MLP模型，接收传感器数据进行推理并下发动作序列；")
add_para(doc, "（c）通信协议：gRPC（protobuf），图像3×512×512 JPEG压缩，力/触觉float32数组，动作50步×6D，总延迟≤220ms/帧。")

add_heading_styled(doc, "3. 安全功能分层部署", level=3)
t_sf = doc.add_table(rows=5, cols=4, style='Table Grid')
for i, hdr in enumerate(["功能", "L2基线", "L3增强", "L4旗舰"]):
    c = t_sf.cell(0, i); c.text = hdr
    set_cell_shading(c, '00d4aa')
    for pp in c.paragraphs:
        for r in pp.runs: r.bold = True; r.font.color.rgb = RGBColor(0x06,0x08,0x0d); r.font.size = Pt(9)
sf_data = [
    ["光幕联动安全（靠近→降速→停机）", "✅", "✅", "✅"],
    ["触觉闭环反馈（夹持力>2N释放）", "—", "✅", "✅"],
    ["异常诊断自恢复（卡料/偏移/失败）", "—", "—", "✅"],
    ["五层主动保护", "—", "—", "✅"],
]
for ri, row in enumerate(sf_data):
    for ci, val in enumerate(row):
        c = t_sf.cell(ri+1, ci); c.text = val
        if ri % 2 == 0: set_cell_shading(c, '0d1117')
        else: set_cell_shading(c, '080c14')
        for pp in c.paragraphs:
            for r in pp.runs: r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xc8,0xd1,0xd9)
doc.add_paragraph()

# ===== 七、权利要求 =====
add_heading_styled(doc, "七、权利要求书", level=1)
doc.add_paragraph()
add_heading_styled(doc, "九、本地-云端统一推理框架", level=1)
add_para(doc, "本实用新型进一步提供一种本地-云端统一推理架构，其特征在于：")
add_para(doc, "（1）本地端(Sys-1)运行轻量ACT模型(52M,<10ms推理)，满足实时控制需求；")
add_para(doc, "（2）云端(Sys-2)部署VTLA(450M)和GR00T(2B)大模型，通过HTTP/gRPC接口提供远程推理服务；")
add_para(doc, "（3）Sys-1通过引擎热切换机制，可在本地ACT、远程VTLA、远程GR00T之间无缝切换；")
add_para(doc, "（4）当云端不可达时自动回退至本地ACT引擎，保证系统可靠性。")
add_para(doc, "1. 一种基于多模态视觉-语言-动作模型的具身机器人精细操作控制系统，其特征在于，包括：", bold=True)
add_para(doc, "感知层，用于采集环境多模态信息，包括3D深度相机、腕部视觉模块、六维力/力矩传感器和触觉传感器阵列；")
add_para(doc, "模型层，包括多模态视觉-语言-动作模型，该模型由视觉编码器、语言编码器、动作解码器和内嵌世界模型组成，用于生成精细操作动作序列；")
add_para(doc, "控制层，包括双臂协同控制器和力控闭环控制器，其中力控闭环频率大于1kHz；")
add_para(doc, "执行层，包括轮式移动底盘、双臂力控机械臂和可切换的末端执行器。")
add_para(doc, "2. 根据权利要求1所述的系统，其特征在于，所述多模态视觉-语言-动作模型采用类脑四阶段迭代架构：感知阶段提取多模态特征→理解阶段生成任务嵌入→决策阶段输出动作序列→执行阶段通过力控闭环执行并反馈结果。", bold=True)
add_para(doc, "3. 根据权利要求1所述的系统，其特征在于，所述双臂协同控制器采用异步并行控制策略：左臂执行取料工序、右臂执行插拔工序，通过中转区解耦，整体单次节拍<25秒。", bold=True)
add_para(doc, "4. 根据权利要求1所述的系统，其特征在于，所述力控闭环控制器同时采集关节力矩信号、末端六维力信号和指尖触觉信号，并通过三路冗余校验实现多模态安全监测。", bold=True)
add_para(doc, "5. 根据权利要求1所述的系统，其特征在于，所述系统通过纯软件升级在L2基线版（分段式自动化）、L3增强版（多模态端到端）和L4旗舰版（精细感知全自主）之间切换，无需更换硬件。", bold=True)
add_para(doc, "6. 根据权利要求1所述的系统，其特征在于，所述世界模型用于预测动作执行后的环境状态变化，并在动作执行前进行安全性校验。", bold=True)

# ===== 八、摘要 =====
add_heading_styled(doc, "八、摘要", level=1)
add_para(doc, "本实用新型公开了一种基于多模态视觉-语言-动作（VLA）模型的具身机器人精细操作控制系统。该系统通过感知层采集3D视觉、力觉和触觉等多模态信息，由多模态VLA模型生成精细操作动作序列，经双臂协同控制器和力控闭环控制器驱动机器人执行。系统采用类脑四阶段迭代架构（感知→理解→决策→执行），可达到±0.02mm插拔精度和>99%关键工序良率，且同一硬件平台支持L2→L3→L4纯软件升级。本实用新型适用于光模块、半导体等高精度电子元器件的自动化插拔、装配与测试场景，具有精度高、自适应强、柔性好、安全可靠等优点。")

# 保存
output_path = "/home/xspace/lerobot-smolvla-lew/docs/patents/Z-MAX-专利交底书-实用新型-多模态VLA具身机器人精细操作控制系统.docx"
doc.save(output_path)
print(f"✅ 专利交底书已保存: {output_path}")
print(f"   文件大小: {os.path.getsize(output_path)} bytes")
