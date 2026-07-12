#!/usr/bin/env python3
"""生成 Z-MAX 用户需求调研问卷 (docx)"""
import os, sys
sys.path.insert(0, '/home/xspace/miniconda3/envs/lerobot/lib/python3.12/site-packages')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def h(text, level=1):
    hh = doc.add_heading(text, level=level)
    for r in hh.runs:
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return hh

def p(text, bold=False, size=11, align=None, space_after=6):
    pp = doc.add_paragraph()
    r = pp.add_run(text)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(size)
    r.bold = bold
    if align: pp.alignment = align
    pp.paragraph_format.space_after = Pt(space_after)
    pp.paragraph_format.line_spacing = 1.5
    return pp

def cell_shade(cell, color):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), color)
    sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

def make_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(headers):
        c = t.cell(0, i)
        c.text = hdr
        cell_shade(c, '00d4aa')
        for pp in c.paragraphs:
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in pp.runs:
                r.bold = True; r.font.color.rgb = RGBColor(0x06,0x08,0x0d); r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.cell(ri+1, ci)
            c.text = str(val)
            if ri % 2 == 0: cell_shade(c, '0d1117')
            else: cell_shade(c, '080c14')
            for pp in c.paragraphs:
                for r in pp.runs:
                    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xc8,0xd1,0xd9)
    doc.add_paragraph()
    return t

# ═══ COVER ═══
for _ in range(5): doc.add_paragraph()
p("Z-MAX 多模态动作专家", bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
p("用户需求调研问卷", bold=True, size=28, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
p("—— 助力光模块工厂实现全自主精密操作 ——", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

t1 = doc.add_table(rows=5, cols=2, style='Table Grid')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
info = [("文档编号","ZFCY-SURVEY-2026-001"),("版本","v1.0.4"),("编制","智蜂创元 ZFCY · 产品管理部"),
        ("日期","2026年07月"),("保密","客户机密 · 仅限项目评估使用")]
for i,(k,v) in enumerate(info):
    t1.cell(i,0).text=k; t1.cell(i,1).text=v
    cell_shade(t1.cell(i,0),'E8E8E8')
    for pp in t1.cell(i,0).paragraphs:
        pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in pp.runs: r.bold=True; r.font.size=Pt(10)

doc.add_page_break()

# ═══ 填表说明 ═══
h("填表说明", 1)
p("本问卷旨在全面了解贵单位在光模块测试/装配环节的自动化需求，以便Z-MAX团队提供针对性的解决方案和投资回报分析。")
p("请根据实际情况如实填写，所有信息仅用于项目评估，严格保密。", space_after=12)
p("填写方式：在 □ 处打勾 ✓，在 _____ 处填写具体内容。", bold=True, space_after=12)

# ═══ 一、客户基本信息 ═══
doc.add_page_break()
h("一、客户基本信息", 1)
make_table(["项目","填写内容"], [
    ["公司/单位名称","________________________"],
    ["所属行业","□ 光模块制造  □ 光通信设备  □ 半导体封装  □ 电子代工  □ 其他:_____"],
    ["公司规模","□ <100人  □ 100-500人  □ 500-2000人  □ >2000人"],
    ["联系人 / 职位","________________ / ________________"],
    ["联系电话 / 邮箱","________________ / ________________"],
    ["现有自动化程度","□ 全人工  □ 半自动  □ 单机自动化  □ 产线自动化"],
])

# ═══ 二、产线现状 ═══
h("二、产线现状", 1)
make_table(["项目","填写内容"], [
    ["产线类型","□ EVB测试线  □ 模块组装线  □ 老化测试线  □ 成品检测线  □ 其他:_____"],
    ["日产能（模块数）","□ <1000  □ 1000-5000  □ 5000-20000  □ >20000"],
    ["班次制度","□ 单班(8h)  □ 双班(16h)  □ 三班(24h)"],
    ["产线操作工人数","_____ 人/班"],
    ["当前关键工序良率","□ <95%  □ 95%-97%  □ 97%-99%  □ >99%"],
    ["主要痛点\n(可多选)","□ 插拔力度不可控  □ 换型耗时长  □ 测试仓空间狭小\n□ 人工误操作多  □ 招工困难  □ 良率波动大\n□ 其他:_____"],
])

# ═══ 三、产品规格 ═══
h("三、产品规格", 1)
p("请列出产线上需要操作的光模块型号：", bold=True, space_after=8)
make_table(["序号","光模块型号","封装类型","传输速率","日产量","操作工序"], [
    ["1","________","□ QSFP-DD □ OSFP □ SFP □ CFP2 □ 其他","□ 100G □ 400G □ 800G □ 1.6T","____个/天","□ 插拔测试 □ 组装 □ AOI □ 老化 □ 其他"],
    ["2","________","□ QSFP-DD □ OSFP □ SFP □ CFP2 □ 其他","□ 100G □ 400G □ 800G □ 1.6T","____个/天","□ 插拔测试 □ 组装 □ AOI □ 老化 □ 其他"],
    ["3","________","□ QSFP-DD □ OSFP □ SFP □ CFP2 □ 其他","□ 100G □ 400G □ 800G □ 1.6T","____个/天","□ 插拔测试 □ 组装 □ AOI □ 老化 □ 其他"],
])

p("需要同时兼容的型号种类数：_____ 种", bold=True)
p("不同型号之间是否需要频繁切换：□ 是（频次:____次/天） □ 否", bold=True)

# ═══ 四、技术要求 ═══
doc.add_page_break()
h("四、技术要求", 1)
make_table(["指标","需求值","备注"], [
    ["插拔对准精度","□ ±0.1mm  □ ±0.05mm  □ ±0.02mm  □ ±0.01mm","____"],
    ["单次操作节拍","□ <10s  □ <8s  □ <25s  □ <3s","____"],
    ["关键工序良率目标","□ ≥97%  □ ≥99%  □ ≥99.5%  □ ≥99.9%","____"],
    ["力控保护要求","□ 需要（最大插入力:____N） □ 不需要","____"],
    ["视觉定位需求","□ 3D视觉引导  □ 2D视觉引导  □ 不需要","____"],
    ["扫码追溯","□ 需要（条码类型:____） □ 不需要","____"],
    ["AOI检测","□ 需要（检测项:____） □ 不需要","____"],
    ["连续工作时长","□ 8h  □ 16h  □ 24h(7×24)","____"],
    ["环境洁净度","□ 百级  □ 千级  □ 万级  □ 无要求","____"],
    ["ESD防护","□ 需要  □ 不需要","____"],
])

# ═══ 五、场地条件 ═══
h("五、场地条件", 1)
make_table(["项目","填写内容"], [
    ["可用场地面积","长____m × 宽____m = ____㎡"],
    ["地面平整度","□ 环氧地坪  □ PVC地板  □ 水泥地  □ 其他:_____"],
    ["供电","□ 220V单相  □ 380V三相  可用功率:____kW"],
    ["压缩空气","□ 有（压力:____MPa） □ 无"],
    ["网络","□ 以太网  □ WiFi  □ 无"],
    ["产线对接信号","□ PLC（型号:____） □ MES  □ 无"],
    ["安全要求","□ 安全光栅  □ 防护围栏  □ 急停按钮  □ 其他:_____"],
])

# ═══ 六、投资预算 ═══
h("六、投资预算", 1)
make_table(["项目","填写内容"], [
    ["预算范围","□ <50万  □ 50-100万  □ 100-300万  □ 300万+"],
    ["预期ROI回收期","□ <12月  □ 12-24月  □ 24-36月  □ 无要求"],
    ["采购计划","□ 1台试点  □ 2-5台  □ 5-10台  □ 10台以上"],
    ["决策周期","□ 1个月内  □ 1-3个月  □ 3-6个月  □ 6个月以上"],
])

# ═══ 七、其他需求 ═══
h("七、其他需求与期望", 1)
p("7.1 您对自动化方案的期望和关注点：", bold=True)
p("_____________________________________________________________________________")
p("_____________________________________________________________________________")
p("_____________________________________________________________________________")
p("7.2 是否有现有自动化方案（如有，请简要说明）：", bold=True)
p("_____________________________________________________________________________")
p("_____________________________________________________________________________")
p("7.3 是否有意向参观Z-MAX演示现场：", bold=True)
p("□ 是（期望时间:__________） □ 暂不需要")
p("7.4 其他补充说明：", bold=True)
p("_____________________________________________________________________________")
p("_____________________________________________________________________________")

# ═══ 底部 ═══
doc.add_page_break()
h("联系方式", 1)
p("感谢您抽出宝贵时间填写本问卷。问卷提交后，Z-MAX团队将在3个工作日内与您联系，提供定制化解决方案和ROI分析。", space_after=16)

make_table(["渠道","信息"], [
    ["公司","智蜂创元 ZFCY"],
    ["产品","Z-MAX 多模态动作专家 (v1.0.4)"],
    ["网站","http://datadrive.world"],
    ["邮箱","（待填写）"],
    ["电话","（待填写）"],
    ["GitHub","https://github.com/MikeBMW/lerobot-smolvla-lew"],
])

p("")
p("Z-MAX —— 以智能重新定义精密制造", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

# 保存
os.makedirs("/home/xspace/lerobot-smolvla-lew/docs/survey", exist_ok=True)
out = "/home/xspace/lerobot-smolvla-lew/docs/survey/Z-MAX-用户需求调研问卷-v1.0.4.docx"
doc.save(out)
print(f"✅ {out} ({os.path.getsize(out)} bytes)")
